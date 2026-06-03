import re
import os
import requests
import json


def ollama_generate(prompt, model="qwen2.5:7b"):
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": model, "prompt": prompt},
            stream=True
        )

        full_text = ""

        for line in response.iter_lines():
            if not line:
                continue

            try:
                data = json.loads(line.decode("utf-8"))
                chunk = data.get("response", "")
                full_text += chunk
            except Exception as e:
                print("🔥 STREAM PARSE ERROR:", e)

        return full_text

    except Exception as e:
        return f"⚠️ Ошибка подключения к Ollama: {e}"

from utils.table_profile_loader import TableProfileLoader
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

class FMFHIFormatChecker:
    """Проверка формата документа по стандартам ФМФХИ"""

    # -----------------------------
    # ПАТТЕРНЫ РАСПОЗНАВАНИЯ
    # -----------------------------

    CHAPTER_PATTERN = re.compile(
        r"^глава\s*№?\s*\d+[\.\:\-\s]*.*$",
        re.IGNORECASE
    )

    SUBCHAPTER_PATTERN = re.compile(
        r"^\s*(?:§\s*)?\d+\.\d+(?:\.\d+)?(?:\s|\.|:|-|$)",
        re.IGNORECASE
    )

    APPENDIX_PATTERN = re.compile(
        r"^приложени[ея]\s*(?:№?\s*[а-яa-z0-9])?\s*$",
        re.IGNORECASE
    )

    FIGURE_PATTERN = re.compile(
        r"^(?:рис\.?\s*\d+|рисунок\s+\d+|figure\s+\d+|fig\.?\s*\d+)",
        re.IGNORECASE
    )

    LITERATURE_VARIANTS = [
        "литература",
        "список литературы",
        "список использованной литературы",
        "список используемой литературы",
        "список использованных источников",
        "список используемых источников",
        "библиографический список",
        "библиография",
        "использованная литература",
        "использованные источники",
        "список источников",
        "список литературы и источников",
    ]

    # -----------------------------
    # ИНИЦИАЛИЗАЦИЯ
    # -----------------------------

    def __init__(self, doc, profile):
        self.doc = doc
        self.profile = profile
        self.errors_found = 0
        self.table_profile = TableProfileLoader.load_table_profile()

    # -----------------------------
    # РАСПОЗНАВАНИЕ ЭЛЕМЕНТОВ
    # -----------------------------

    def is_main_heading(self, text: str) -> bool:
        t = text.strip().lower()
        if t in {"введение", "заключение"}:
            return True
        if t in self.LITERATURE_VARIANTS:
            return True
        return bool(self.CHAPTER_PATTERN.match(t))

    def is_subheading(self, text: str) -> bool:
        return bool(self.SUBCHAPTER_PATTERN.match(text.strip()))

    def is_appendix(self, text: str) -> bool:
        return bool(self.APPENDIX_PATTERN.match(text.strip()))

    def is_literature_heading(self, text: str) -> bool:
        t = text.strip().lower()
        if t in self.LITERATURE_VARIANTS:
            return True
        return any(k in t for k in ["литератур", "источник", "библиограф"])

    def is_figure_caption(self, text: str) -> bool:
        return bool(self.FIGURE_PATTERN.match(text.strip()))

    def is_equation(self, paragraph) -> bool:
    # Формулы Word (OMML)
        if paragraph._p.xpath(".//m:oMath") or paragraph._p.xpath(".//m:oMathPara"):
            return True

        if paragraph._p.xpath(".//w:oMath") or paragraph._p.xpath(".//w:oMathPara"):
            return True

        # Иногда формулы попадают внутрь run
        for run in paragraph.runs:
            if run._r.xpath(".//m:oMath") or run._r.xpath(".//m:oMathPara"):
                return True
            if run._r.xpath(".//w:oMath") or run._r.xpath(".//w:oMathPara"):
                return True

        return False


    def is_table_content(self, paragraph) -> bool:
        parent = paragraph._p.getparent()
        while parent is not None:
            if parent.tag.endswith('tc'):
                return True
            parent = parent.getparent()
        return False

    # -----------------------------
    # ПОИСК ВВЕДЕНИЯ
    # -----------------------------

    def find_intro_index(self) -> int:
        for i, p in enumerate(self.doc.paragraphs):
            text = p.text.strip().lower()

            # Убираем мусор
            clean = text.replace(":", "").replace(".", "").replace("–", "").strip()

            # Убираем двойные пробелы
            clean = " ".join(clean.split())

            # Убираем неразрывные пробелы
            clean = clean.replace("\u00A0", " ")

            # Проверяем варианты
            if clean in ("введение", "введениe", "введенне", "введенiе"):
                return i

        return 0


    # -----------------------------
    # ПАРАМЕТРЫ ПАРАГРАФА
    # -----------------------------

    def get_paragraph_font(self, paragraph) -> str:
        fonts_found = []
        for run in paragraph.runs:
            if run.font.name:
                fonts_found.append(run.font.name)
            elif run._r.xpath(".//w:rFonts"):
                rFonts = run._r.xpath(".//w:rFonts")[0]
                ascii_font = rFonts.get(qn('w:ascii'))
                if ascii_font:
                    fonts_found.append(ascii_font)

        if fonts_found:
            return max(set(fonts_found), key=fonts_found.count)

        if paragraph.style and paragraph.style.font and paragraph.style.font.name:
            return paragraph.style.font.name

        normal = self.doc.styles["Normal"]
        return normal.font.name if normal.font.name else None

    def get_paragraph_font_size(self, paragraph):
        sizes = []
        for run in paragraph.runs:
            if run.font.size:
                sizes.append(run.font.size.pt)

        if sizes:
            return min(sizes)

        if paragraph.style and paragraph.style.font and paragraph.style.font.size:
            return paragraph.style.font.size.pt

        normal = self.doc.styles["Normal"]
        return normal.font.size.pt if normal.font.size else None

    def get_alignment(self, paragraph):
        if paragraph.alignment is not None:
            return paragraph.alignment
        if paragraph.style and paragraph.style.paragraph_format.alignment:
            return paragraph.style.paragraph_format.alignment
        return WD_ALIGN_PARAGRAPH.LEFT

    def get_line_spacing(self, paragraph):
        pf = paragraph.paragraph_format
        if pf and pf.line_spacing:
            try:
                return float(pf.line_spacing)
            except:
                pass

        if paragraph.style and paragraph.style.paragraph_format.line_spacing:
            try:
                return float(paragraph.style.paragraph_format.line_spacing)
            except:
                pass

        return None

    # -----------------------------
    # ПОЛЯ ДОКУМЕНТА
    # -----------------------------

    def get_page_margins(self):
        section = self.doc.sections[0]

        def twips_to_mm(val):
            # Length.twips — корректный twips
            if hasattr(val, "twips"):
                twips = val.twips
            else:
                twips = int(val)
            return twips * 25.4 / 1440

        return {
            'left': twips_to_mm(section.left_margin),
            'right': twips_to_mm(section.right_margin),
            'top': twips_to_mm(section.top_margin),
            'bottom': twips_to_mm(section.bottom_margin),
        }

    def check_margins(self):
        errors = []
        margins = self.get_page_margins()

        expected = {'left': 30, 'right': 15, 'top': 20, 'bottom': 20}
        tolerance = 2

        names = {
            'left': 'Левое',
            'right': 'Правое',
            'top': 'Верхнее',
            'bottom': 'Нижнее'
        }

        for side, exp in expected.items():
            actual = round(margins[side], 1)
            if abs(actual - exp) > tolerance:
                errors.append(
                    f"Найдена ошибка. {names[side]} поле: {actual} мм "
                    f"(ожидалось {exp} мм, допуск ±{tolerance} мм)"
                )

        return errors

    # -----------------------------
    # ПРОВЕРКА ТАБЛИЦ
    # -----------------------------

    def check_tables(self):
        results = {}

        expected_font = self.table_profile.get("font", "Times New Roman")
        expected_sizes = self.table_profile.get("sizes", [12, 14])
        expected_align = self.table_profile.get("alignment", "justify")
        expected_ls = float(self.table_profile.get("line_spacing", 1.5))

        align_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }

        human = {
            "left": "по левому краю",
            "right": "по правому краю",
            "center": "по центру",
            "justify": "по ширине",
            "unknown": "неизвестно",
        }

        for idx, table in enumerate(self.doc.tables, start=1):
            name = f"Таблица {idx}"
            errs = []

            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = p.text.strip()
                        if not text:
                            continue

                        preview = text[:50] + ("..." if len(text) > 50 else "")

                        font = self.get_paragraph_font(p)
                        if font and font.lower() != expected_font.lower():
                            errs.append(
                                f"Найдена ошибка в '{preview}' — шрифт: найдено '{font}', ожидалось '{expected_font}'"
                            )

                        size = self.get_paragraph_font_size(p)
                        if size and round(size) not in expected_sizes:
                            errs.append(
                                f"Найдена ошибка в '{preview}' — размер шрифта: {size} pt (ожидалось {expected_sizes})"
                            )

                        ls = self.get_line_spacing(p)
                        if ls is not None and round(ls, 2) != round(expected_ls, 2):
                            errs.append(
                                f"Найдена ошибка в '{preview}' — межстрочный интервал: {ls} (ожидалось {expected_ls})"
                            )

                        align = self.get_alignment(p)
                        align_key = align_map.get(align, "unknown")
                        if align_key != expected_align:
                            errs.append(
                                f"Найдена ошибка в '{preview}' — выравнивание: {human.get(align_key)} "
                                f"(ожидалось {human.get(expected_align)})"
                            )

            if errs:
                results[name] = errs

        return results

    # -----------------------------
    # ПРОВЕРКА ТЕКСТА
    # -----------------------------

    def check_formatting(self) -> list:
        errors = []
        intro_index = self.find_intro_index()

        # Разрешённые шрифты
        ALLOWED_FONTS = {"Times New Roman", "Cambria Math"}

                    # Хак: если строка похожа на формулу, не трогаем её
                    # Хак: если строка похожа на формулу, не трогаем её

        for idx, p in enumerate(self.doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue

            # Пропускаем всё до "Введение"
            if idx < intro_index:
                continue

            # Пропускаем формулы Word (OMML)
            if self.is_equation(p):
                continue

            # Пропускаем текст внутри таблиц
            if self.is_table_content(p):
                continue

            math_like_tokens = ["=", "≤", "≥", "sin", "cos", "tg", "ctg", "│", "|", "√"]
            if any(tok in text for tok in math_like_tokens):
                # Разрешаем таким строкам жить как есть (часто центрированные)
                continue

            preview = text[:80] + ("..." if len(text) > 80 else "")

            # -----------------------------
            # ЗАГОЛОВКИ
            # -----------------------------
            if self.is_main_heading(text):
                if self.get_alignment(p) != WD_ALIGN_PARAGRAPH.CENTER:
                    errors.append(
                        f"Найдена ошибка. \"{preview}\" — заголовок должен быть выровнен по центру"
                    )
                continue

            if self.is_subheading(text):
                if self.get_alignment(p) != WD_ALIGN_PARAGRAPH.CENTER:
                    errors.append(
                        f"Найдена ошибка. \"{preview}\" — подглава должна быть выровнена по центру"
                    )
                continue

            if self.is_appendix(text):
                if self.get_alignment(p) != WD_ALIGN_PARAGRAPH.CENTER:
                    errors.append(
                        f"Найдена ошибка. \"{preview}\" — слово \"Приложение\" должно быть выровнено по центру"
                    )
                continue

            if self.is_literature_heading(text):
                if self.get_alignment(p) != WD_ALIGN_PARAGRAPH.CENTER:
                    errors.append(
                        f"Найдена ошибка. \"{preview}\" — заголовок списка литературы должен быть выровнен по центру"
                    )
                continue

            if self.is_figure_caption(text):
                if self.get_alignment(p) != WD_ALIGN_PARAGRAPH.CENTER:
                    errors.append(
                        f"Найдена ошибка. \"{preview}\" — подпись к рисунку должна быть выровнена по центру"
                    )
                continue

            # -----------------------------
            # Обычный текст
            # -----------------------------
            style_name = p.style.name if p.style else "Normal"

            # Если стиль заголовка, но текст НЕ заголовок → считаем обычным текстом
            if style_name in ("Heading 1", "Heading 2") and not (
                self.is_main_heading(text) or self.is_subheading(text)
            ):
                style_name = "Normal"

            # Если стиль не описан в профиле — пропускаем
            if style_name not in self.profile:
                continue

            expected = self.profile[style_name]

            # -----------------------------
            # ШРИФТ
            # -----------------------------
            font = self.get_paragraph_font(p)
            if font and font not in ALLOWED_FONTS:
                errors.append(
                    f"Найдена ошибка. \"{preview}\" — шрифт: найдено '{font}', "
                    f"ожидалось один из {ALLOWED_FONTS}"
                )

            # -----------------------------
            # РАЗМЕР ШРИФТА
            # -----------------------------
            size = self.get_paragraph_font_size(p)
            if size and round(size) != expected["size"]:
                errors.append(
                    f"Найдена ошибка. \"{preview}\" — размер шрифта: {size} pt "
                    f"(ожидалось {expected['size']} pt)"
                )

            # -----------------------------
            # ВЫРАВНИВАНИЕ
            # -----------------------------
            align = self.get_alignment(p)
            if align != WD_ALIGN_PARAGRAPH.JUSTIFY:
                align_map = {
                    WD_ALIGN_PARAGRAPH.LEFT: "по левому краю",
                    WD_ALIGN_PARAGRAPH.RIGHT: "по правому краю",
                    WD_ALIGN_PARAGRAPH.CENTER: "по центру",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "по ширине",
                }
                found_align = align_map.get(align, "неизвестно")
                errors.append(
                    f"Найдена ошибка. \"{preview}\" — выравнивание: {found_align} "
                    f"(ожидалось по ширине)"
                )

            # -----------------------------
            # МЕЖСТРОЧНЫЙ ИНТЕРВАЛ
            # -----------------------------
            ls = self.get_line_spacing(p)
            if ls is not None and round(ls, 2) != 1.5:
                errors.append(
                    f"Найдена ошибка. \"{preview}\" — межстрочный интервал: {ls} "
                    f"(ожидалось 1.5)"
                )

        return errors


    # -----------------------------
    # ФОРМАТИРОВАНИЕ ОТЧЁТА
    # -----------------------------

    def format_error_block(self, title: str, errors: list, emoji: str = "📌") -> list:
        if not errors:
            return []

        output = []
        header = f"{emoji} {title} ({len(errors)} ошибок)"
        sep = "=" * (len(header) + 4)

        output.append(sep)
        output.append(f"  {header}")
        output.append(sep)

        for i, err in enumerate(errors, start=1):
            output.append(f"[{i:02d}] {err}")
            if i < len(errors):
                output.append("─" * 60)

        output.append("=" * 60)
        output.append("")
        return output

    # -----------------------------
    # СВОДКА
    # -----------------------------

    def create_summary(self, formatting_errors, table_errors, margin_errors):
        total = len(formatting_errors) + sum(len(v) for v in table_errors.values()) + len(margin_errors)
        self.errors_found = total

        summary = [
            "📘 ИНСТРУКЦИЯ ПО ОФОРМЛЕНИЮ",
            "────────────────────────────────────────────────────────────",
            "🔹 Главы и подглавы должны быть выровнены по центру.",
            "🔹 Подглавы (1.1, 1.2, 2.3…) — это заголовки, а не списки.",
            "🔹 Поля: левое – 30 мм, правое – 15 мм, верхнее – 20 мм, нижнее – 20 мм.",
            "────────────────────────────────────────────────────────────",
            "",
            "╔══════════════════════════════════════════════════════════╗",
            "║                    РЕЗУЛЬТАТЫ ПРОВЕРКИ                   ║",
            "╠══════════════════════════════════════════════════════════╣",
            f"║ 📄 Ошибки полей:           {len(margin_errors):3d}                ║",
            f"║ ⚠️ Ошибки текста:          {len(formatting_errors):3d}                ║",
            f"║ 📋 Ошибки таблиц:          {sum(len(v) for v in table_errors.values()):3d}                ║",
            "╠══════════════════════════════════════════════════════════╣",
            f"║ 🔴 Всего ошибок:           {total:3d}                ║",
            "╚══════════════════════════════════════════════════════════╝",
            "",
        ]

        if total == 0:
            summary.append("🎉 Документ полностью соответствует требованиям ФМФХИ!")
            summary.append("")

        return summary

    # -----------------------------
    # ЭКСПОРТ ОТЧЁТА
    # -----------------------------

    def save_report_as_txt(self, report_lines, file_path):
        with open(file_path, "w", encoding="utf-8") as f:
            for line in report_lines:
                f.write(line + "\n")

    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet

    def save_report_as_pdf(self, report_lines: list, file_path: str):

        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import os
        import re

        # --- Путь к шрифту ---
        font_path = os.path.join(os.path.dirname(__file__), "DejaVuSans.ttf")

        if not os.path.exists(font_path):
            raise FileNotFoundError(f"Не найден шрифт: {font_path}")

        # --- Регистрируем шрифт ---
        pdfmetrics.registerFont(TTFont('DejaVu', font_path))

        # --- Документ ---
        doc = SimpleDocTemplate(file_path)
        styles = getSampleStyleSheet()

        # --- Стили ---
        normal_style = ParagraphStyle(
            name='NormalDejaVu',
            parent=styles['Normal'],
            fontName='DejaVu',
            fontSize=12,
            leading=14
        )

        heading_style = ParagraphStyle(
            name='HeadingDejaVu',
            parent=styles['Heading1'],
            fontName='DejaVu',
            fontSize=16,
            leading=18
        )

        subheading_style = ParagraphStyle(
            name='SubHeadingDejaVu',
            parent=styles['Heading2'],
            fontName='DejaVu',
            fontSize=14,
            leading=16
        )

        # --- Удаление эмодзи ---
        def remove_emoji(text):
            return re.sub(r'[^\w\s.,:;!?()\-\[\]"/]', '', text)

        elements = []

        for line in report_lines:
            if not line.strip():
                elements.append(Spacer(1, 10))
                continue

            # Убираем эмодзи
            clean_line = remove_emoji(line)

            # Экранируем спецсимволы
            safe_line = (
                clean_line.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
            )

            # --- Логика оформления ---
            if "ИНСТРУКЦИЯ" in line:
                elements.append(Paragraph(f"<b>{safe_line}</b>", heading_style))

            elif "ОШИБКИ" in line or "Таблица" in line:
                elements.append(Paragraph(f"<b>{safe_line}</b>", subheading_style))

            else:
                elements.append(Paragraph(safe_line, normal_style))

            elements.append(Spacer(1, 6))

        # --- Сборка PDF ---
        doc.build(elements)

    # -----------------------------
    # ОСНОВНОЙ МЕТОД ПРОВЕРКИ
    # -----------------------------

    def run_checks(self) -> list:
        margin_errors = self.check_margins()
        formatting_errors = self.check_formatting()
        table_errors = self.check_tables()

        report = []
        report.extend(self.create_summary(formatting_errors, table_errors, margin_errors))

        if margin_errors:
            report.extend(self.format_error_block("ОШИБКИ ПОЛЕЙ", margin_errors, emoji="📐"))

        if formatting_errors:
            report.extend(self.format_error_block("ОШИБКИ ТЕКСТА", formatting_errors, emoji="⚠️"))

        for table_name, errs in table_errors.items():
            report.extend(self.format_error_block(f"ОШИБКИ В {table_name}", errs, emoji="📊"))

        # -----------------------------
        # ДОБАВЛЯЕМ ИИ‑ПРОВЕРКУ
        # -----------------------------
        ai_report = self.run_ai_check()

        report.append("")
        report.append("🤖 ИИ‑АНАЛИЗ ДОКУМЕНТА")
        report.append("────────────────────────────────────────────")
        report.append(ai_report)
        report.append("")

        return report
    
    def run_ai_check(self):
        # Собираем текст документа
        full_text = "\n".join(p.text for p in self.doc.paragraphs)

        # Формируем промпт
        prompt = f"""
    Ты — эксперт по оформлению курсовых работ ФМФХИ.

    Проанализируй документ по следующим критериям:

    1. Соответствие структуре:
    - наличие введения
    - наличие глав и подглав
    - наличие заключения
    - наличие списка литературы
    - наличие приложений
    - общий объём (норма: 25–30 страниц)

    2. Соответствие требованиям оформления:
    - поля: левое 30 мм, правое 15 мм, верх/низ 20 мм
    - шрифт: Times New Roman, 14 pt
    - межстрочный интервал: 1.5
    - выравнивание: по ширине
    - красная строка: 1.25 см
    - автоматические переносы

    3. Дай рекомендации по улучшению.

    Ответ должен быть кратким, структурированным, в формате Markdown.

    Текст документа:
    ----------------
    {full_text}
    """

        # Логи
        print("🔥 OLLAMA PROMPT ОТПРАВЛЁН")
        print(prompt[:500], "...")

        # Вызов модели
        ai_report = ollama_generate(prompt)

        print("🔥 OLLAMA RESPONSE:")
        print(ai_report[:500], "...")

        return ai_report

                               

